#!/usr/bin/env python3
"""
TEMPLATE BANK SOAL CBT SCHOOL — v3 (Tanpa Header, 5 Jenis Soal, Bergambar)
Generator: python-docx + Pillow
"""

import sys, os, io

SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
PUBLIC_DIR  = os.path.normpath(os.path.join(SCRIPT_DIR, '..', 'frontend', 'public'))
OUTPUT_PATH = os.path.join(PUBLIC_DIR, 'TEMPLATE_SOAL_CBT.docx')
os.makedirs(PUBLIC_DIR, exist_ok=True)

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

# ── Color Palette ──────────────────────────────────────────────────────────────
def hx(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

NAVY  = hx('1B2B4B')
BLUE  = hx('2563EB')
DGRAY = hx('374151')
WHITE = hx('FFFFFF')
GREEN = hx('15803D')
AMBER = hx('B45309')

# Warna per opsi (bg_hex, text_hex)
OPSI_CLR = {
    'A': ('DBEAFE', '1E3A8A'),
    'B': ('DCFCE7', '14532D'),
    'C': ('FEF9C3', '713F12'),
    'D': ('F3E8FF', '581C87'),
    'E': ('FFE4E6', '881337'),
}
CORRECT_BG = 'BBF7D0'
CORRECT_FG = '166534'
HEADER_BG  = '1E3A8A'
HEADER_FG  = 'FFFFFF'
ALT_ROW    = 'F8FAFC'
SOAL_BG    = 'EFF6FF'

# ── Font loader ────────────────────────────────────────────────────────────────
FONT_BOLD = FONT_REG = None
for path in [
    '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
    '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
    '/usr/share/fonts/truetype/freefont/FreeSansBold.ttf',
]:
    if os.path.exists(path):
        try:
            FONT_BOLD = ImageFont.truetype(path, 14)
            reg = path.replace('Bold', '')
            FONT_REG = ImageFont.truetype(reg, 11) if os.path.exists(reg) else FONT_BOLD
            FONT_SM  = ImageFont.truetype(reg, 10) if os.path.exists(reg) else FONT_BOLD
            break
        except: pass
if FONT_BOLD is None:
    FONT_BOLD = FONT_REG = FONT_SM = ImageFont.load_default()

# ── OOXML Helpers ──────────────────────────────────────────────────────────────
def _el(tag): return OxmlElement(tag)

def shd(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    e = _el('w:shd')
    e.set(qn('w:val'), 'clear'); e.set(qn('w:color'), 'auto')
    e.set(qn('w:fill'), fill_hex.upper()); tcPr.append(e)

def valign(cell, v='top'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for x in tcPr.findall(qn('w:vAlign')): tcPr.remove(x)
    e = _el('w:vAlign'); e.set(qn('w:val'), v); tcPr.append(e)

def cell_width(cell, dxa):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for x in tcPr.findall(qn('w:tcW')): tcPr.remove(x)
    e = _el('w:tcW'); e.set(qn('w:w'), str(dxa)); e.set(qn('w:type'), 'dxa')
    tcPr.append(e)

def cell_mar(cell, top=60, bot=60, left=100, right=80):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for x in tcPr.findall(qn('w:tcMar')): tcPr.remove(x)
    m = _el('w:tcMar')
    for side, val in [('top',top),('bottom',bot),('left',left),('right',right)]:
        e = _el(f'w:{side}'); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)

def pspacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for x in pPr.findall(qn('w:spacing')): pPr.remove(x)
    e = _el('w:spacing'); e.set(qn('w:before'), str(int(before))); e.set(qn('w:after'), str(int(after)))
    pPr.append(e)

def run_text(para, text, font='Calibri', size=9, bold=False, italic=False, color=None):
    r = para.add_run(text)
    r.font.name = font; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return r

def write_cell(cell, text, font='Calibri', size=9, bold=False, italic=False,
               color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
               bg=None, v='top', mar=(60,60,100,80)):
    if bg: shd(cell, bg)
    valign(cell, v); cell_mar(cell, *mar)
    for p in cell.paragraphs[1:]: p._p.getparent().remove(p._p)
    para = cell.paragraphs[0]; para.clear(); para.alignment = align
    pspacing(para, 0, 0)
    if text: run_text(para, text, font, size, bold, italic, color)
    return para

def add_para(cell, text=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             font='Calibri', size=9, bold=False, color=None):
    para = cell.add_paragraph(); para.alignment = align; pspacing(para, 0, 0)
    if text: run_text(para, text, font, size, bold, color=color)
    return para

def set_borders(table, color='CCCCCC', size=6):
    tbl = table._tbl; tblPr = tbl.tblPr
    for x in tblPr.findall(qn('w:tblBorders')): tblPr.remove(x)
    borders = _el('w:tblBorders')
    for s in ['top','left','bottom','right','insideH','insideV']:
        b = _el(f'w:{s}'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(size))
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),color); borders.append(b)
    tblPr.append(borders)

def tbl_width(table, dxa):
    tbl = table._tbl; tblPr = tbl.tblPr
    for x in tblPr.findall(qn('w:tblW')): tblPr.remove(x)
    e = _el('w:tblW'); e.set(qn('w:w'), str(dxa)); e.set(qn('w:type'), 'dxa')
    tblPr.append(e)

def hdr_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    e = _el('w:tblHeader'); trPr.append(e)

def no_border_table(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    for x in tblPr.findall(qn('w:tblBorders')): tblPr.remove(x)
    borders = _el('w:tblBorders')
    for s in ['top','left','bottom','right','insideH','insideV']:
        b = _el(f'w:{s}'); b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto'); borders.append(b)
    tblPr.append(borders)

# ── Image Generators ───────────────────────────────────────────────────────────
def make_photo_placeholder(w, h, label='FOTO / GAMBAR SOAL'):
    """Gambar placeholder realistis dengan ikon kamera."""
    img = Image.new('RGB', (w, h), (230, 240, 255))
    draw = ImageDraw.Draw(img)
    # Background gradient effect
    for y in range(h):
        ratio = y / h
        r = int(230 + (210-230)*ratio)
        g = int(240 + (225-240)*ratio)
        b = int(255 + (245-255)*ratio)
        draw.line([(0,y),(w,y)], fill=(r,g,b))
    # Border dashed
    for x in range(0, w, 14):
        draw.line([(x,4),(min(x+8,w),4)], fill=(100,140,220), width=2)
        draw.line([(x,h-5),(min(x+8,w),h-5)], fill=(100,140,220), width=2)
    for y in range(0, h, 14):
        draw.line([(4,y),(4,min(y+8,h))], fill=(100,140,220), width=2)
        draw.line([(w-5,y),(w-5,min(y+8,h))], fill=(100,140,220), width=2)
    # Camera body
    cx, cy = w//2, h//2 - 14
    cw, ch = min(60,w//2), min(42, h//3)
    draw.rounded_rectangle([cx-cw//2, cy-ch//2, cx+cw//2, cy+ch//2],
                            radius=5, fill=(70,110,200), outline=(40,80,160), width=2)
    # Lens
    lr = min(14, ch//2 - 5)
    draw.ellipse([cx-lr-3, cy-lr-3, cx+lr+3, cy+lr+3], fill=(200,220,255), outline=(40,80,160), width=2)
    draw.ellipse([cx-lr+3, cy-lr+3, cx+lr-3, cy+lr-3], fill=(130,170,240))
    draw.ellipse([cx-lr+6, cy-lr+6, cx+lr-6, cy+lr-6], fill=(70,110,200))
    # Flash
    draw.ellipse([cx+cw//2-14, cy-ch//2+5, cx+cw//2-5, cy-ch//2+14],
                 fill=(255,230,80), outline=(200,160,0), width=1)
    # Label text
    try:
        bb = draw.textbbox((0,0), label, font=FONT_REG)
        tw = bb[2]-bb[0]
        draw.text((cx - tw//2, cy + ch//2 + 10), label, fill=(50,90,180), font=FONT_REG)
    except: pass
    buf = io.BytesIO(); img.save(buf, 'PNG'); buf.seek(0)
    return buf

def make_type_badge(label, bg_hex, fg_hex, w=180, h=32):
    bg = tuple(int(bg_hex[i:i+2],16) for i in (0,2,4))
    fg = tuple(int(fg_hex[i:i+2],16) for i in (0,2,4))
    img = Image.new('RGB', (w, h), (255,255,255))
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle([1,1,w-2,h-2], radius=8, fill=bg, outline=fg, width=2)
    try:
        bb = draw.textbbox((0,0), label, font=FONT_BOLD)
        tw, th = bb[2]-bb[0], bb[3]-bb[1]
        draw.text(((w-tw)//2, (h-th)//2 - bb[1]//2), label, fill=fg, font=FONT_BOLD)
    except: pass
    buf = io.BytesIO(); img.save(buf, 'PNG'); buf.seek(0)
    return buf

# Tabel kolom utama: NO | SOAL | JENIS | OPSI | JAWABAN | KUNCI
# DXA widths (1 cm ≈ 567 dxa). Total ≈ 18cm = 10206 dxa
COL_W = {
    'NO':     400,
    'SOAL':  3000,
    'JENIS':  900,
    'OPSI':   600,
    'JAWABAN':4506,
    'KUNCI':  800,
}
TOTAL_W = sum(COL_W.values())  # 10206

COLS    = list(COL_W.keys())
COL_DXA = list(COL_W.values())

# ── MAIN TABLE HELPERS ─────────────────────────────────────────────────────────
def make_main_table(doc):
    t = doc.add_table(rows=1, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t, '94A3B8', 6)
    tbl_width(t, TOTAL_W)
    return t

def set_row_widths(row):
    for ci, dxa in enumerate(COL_DXA):
        cell_width(row.cells[ci], dxa)

def add_header_row(table):
    row = table.rows[0]
    set_row_widths(row)
    labels = ['NO', 'SOAL / PERTANYAAN', 'JENIS', 'OPSI', 'JAWABAN / PILIHAN', 'KUNCI']
    for ci, lbl in enumerate(labels):
        write_cell(row.cells[ci], lbl, size=9, bold=True,
                   color=hx(HEADER_FG), align=WD_ALIGN_PARAGRAPH.CENTER,
                   bg=HEADER_BG, v='center', mar=(80,80,60,60))
    hdr_repeat(row)

def add_soal_row(table, no, soal_text, jenis_label, jenis_bg, soal_bg=SOAL_BG,
                 with_image=False, image_buf=None):
    """Tambah baris pertama soal (baris SOAL + cell JENIS, NO)."""
    row = table.add_row()
    set_row_widths(row)
    # NO
    write_cell(row.cells[0], str(no), size=9, bold=True,
               color=hx(HEADER_BG), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg='F1F5F9', v='top', mar=(80,60,60,60))
    # SOAL
    sc = row.cells[1]; shd(sc, soal_bg); valign(sc, 'top'); cell_mar(sc, 80, 60, 100, 80)
    for p in sc.paragraphs[1:]: p._p.getparent().remove(p._p)
    p0 = sc.paragraphs[0]; p0.clear(); pspacing(p0, 0, 0)
    run_text(p0, soal_text, 'Calibri', 9.5, False, False, hx('374151'))
    if with_image and image_buf:
        pi = sc.add_paragraph(); pspacing(pi, 60, 0); pi.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pi.add_run().add_picture(image_buf, width=Cm(4.5))
    # JENIS
    jc = row.cells[2]; shd(jc, jenis_bg); valign(jc, 'center'); cell_mar(jc, 80, 80, 60, 60)
    for p in jc.paragraphs[1:]: p._p.getparent().remove(p._p)
    pj = jc.paragraphs[0]; pj.clear(); pj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspacing(pj, 0, 0)
    run_text(pj, jenis_label, 'Calibri', 8, True, False, hx(HEADER_BG))
    # OPSI + JAWABAN + KUNCI — dikosongkan di baris pertama
    for ci in [3, 4, 5]:
        write_cell(row.cells[ci], '', bg=ALT_ROW if no%2==0 else 'FFFFFF')
    return row

def add_option_row(table, no, opsi_key, jawaban_text, kunci_val, is_correct=False):
    """Tambah baris opsi pilihan (OPSI, JAWABAN, KUNCI terisi; NO/SOAL/JENIS kosong)."""
    row = table.add_row()
    set_row_widths(row)
    row_bg = CORRECT_BG if is_correct else (OPSI_CLR.get(opsi_key, ('F8FAFC','374151'))[0])
    txt_fg = CORRECT_FG if is_correct else OPSI_CLR.get(opsi_key, ('F8FAFC','374151'))[1]
    # NO (kosong, warna sama)
    write_cell(row.cells[0], '', bg='F1F5F9')
    # SOAL (kosong)
    write_cell(row.cells[1], '', bg=SOAL_BG if no%2!=0 else ALT_ROW)
    # JENIS (kosong)
    write_cell(row.cells[2], '', bg=OPSI_CLR.get(opsi_key, ('F1F5F9','374151'))[0])
    # OPSI
    write_cell(row.cells[3], opsi_key, size=9, bold=True,
               color=hx(txt_fg), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg=row_bg, v='center', mar=(60,60,40,40))
    # JAWABAN
    write_cell(row.cells[4], jawaban_text, size=9, bold=is_correct,
               color=hx(CORRECT_FG if is_correct else txt_fg),
               bg=row_bg, v='top', mar=(60,60,100,80))
    # KUNCI
    write_cell(row.cells[5], kunci_val, size=10, bold=True,
               color=hx(CORRECT_FG if is_correct else '94A3B8'),
               align=WD_ALIGN_PARAGRAPH.CENTER,
               bg=CORRECT_BG if is_correct else 'F8FAFC', v='center', mar=(60,60,40,40))
    return row

def add_essay_row(table, no, soal_text, jenis_label='ESAI', jenis_bg='FFF3CD',
                  kunci_text='(Jawaban uraian siswa)'):
    """Baris soal esai — tidak ada opsi."""
    row = table.add_row()
    set_row_widths(row)
    write_cell(row.cells[0], str(no), size=9, bold=True,
               color=hx(HEADER_BG), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg='F1F5F9', v='top', mar=(80,60,60,60))
    sc = row.cells[1]; shd(sc, SOAL_BG); valign(sc, 'top'); cell_mar(sc, 80, 60, 100, 80)
    for p in sc.paragraphs[1:]: p._p.getparent().remove(p._p)
    p0 = sc.paragraphs[0]; p0.clear(); pspacing(p0, 0, 0)
    run_text(p0, soal_text, 'Calibri', 9.5)
    jc = row.cells[2]; shd(jc, jenis_bg); valign(jc, 'center'); cell_mar(jc, 80, 80, 60, 60)
    for p in jc.paragraphs[1:]: p._p.getparent().remove(p._p)
    pj = jc.paragraphs[0]; pj.clear(); pj.alignment = WD_ALIGN_PARAGRAPH.CENTER; pspacing(pj, 0, 0)
    run_text(pj, jenis_label, 'Calibri', 8, True, False, hx(HEADER_BG))
    # OPSI (dash)
    write_cell(row.cells[3], '—', size=9, color=hx('94A3B8'),
               align=WD_ALIGN_PARAGRAPH.CENTER, bg='F8FAFC', v='center')
    # JAWABAN (rubrik singkat)
    ac = row.cells[4]; shd(ac, 'FFFDE7'); valign(ac, 'top'); cell_mar(ac, 80, 60, 100, 80)
    for p in ac.paragraphs[1:]: p._p.getparent().remove(p._p)
    pa = ac.paragraphs[0]; pa.clear(); pspacing(pa, 0, 0)
    run_text(pa, kunci_text, 'Calibri', 9, False, True, hx('92400E'))
    # KUNCI (rubrik/skor)
    write_cell(row.cells[5], 'Rubrik', size=8, bold=True,
               color=hx('92400E'), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg='FEF9C3', v='center', mar=(60,60,40,40))
    return row

def add_tf_row(table, no, soal_text, kunci_val):
    """Baris soal Benar/Salah."""
    row = table.add_row()
    set_row_widths(row)
    is_b = kunci_val.upper() in ('B', 'BENAR', 'TRUE')
    write_cell(row.cells[0], str(no), size=9, bold=True,
               color=hx(HEADER_BG), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg='F1F5F9', v='top', mar=(80,60,60,60))
    sc = row.cells[1]; shd(sc, SOAL_BG); valign(sc, 'top'); cell_mar(sc, 80, 60, 100, 80)
    for p in sc.paragraphs[1:]: p._p.getparent().remove(p._p)
    p0 = sc.paragraphs[0]; p0.clear(); pspacing(p0, 0, 0)
    run_text(p0, soal_text, 'Calibri', 9.5)
    jc = row.cells[2]; shd(jc, 'D1FAE5'); valign(jc, 'center'); cell_mar(jc, 80, 80, 60, 60)
    for p in jc.paragraphs[1:]: p._p.getparent().remove(p._p)
    pj = jc.paragraphs[0]; pj.clear(); pj.alignment = WD_ALIGN_PARAGRAPH.CENTER; pspacing(pj, 0, 0)
    run_text(pj, 'B/S', 'Calibri', 8, True, False, hx(HEADER_BG))
    # OPSI (B/S)
    write_cell(row.cells[3], 'B/S', size=9, bold=True,
               color=hx('166534' if is_b else '991B1B'),
               align=WD_ALIGN_PARAGRAPH.CENTER,
               bg='D1FAE5' if is_b else 'FEE2E2', v='center')
    # JAWABAN
    ans = 'BENAR' if is_b else 'SALAH'
    write_cell(row.cells[4], ans, size=9, bold=True,
               color=hx('166534' if is_b else '991B1B'),
               bg='D1FAE5' if is_b else 'FEE2E2', v='center')
    # KUNCI
    write_cell(row.cells[5], 'V' if is_b else '', size=10, bold=True,
               color=hx(CORRECT_FG),
               align=WD_ALIGN_PARAGRAPH.CENTER,
               bg=CORRECT_BG if is_b else 'F8FAFC', v='center')
    return row

def add_matching_pair(table, no_label, pasangan_kiri, pasangan_kanan, opsi_key, is_first=False):
    """Baris menjodohkan: OPSI = huruf pasangan, JAWABAN kiri → kanan."""
    row = table.add_row()
    set_row_widths(row)
    row_bg = OPSI_CLR.get(opsi_key, ('F8FAFC','374151'))[0]
    fg_clr = OPSI_CLR.get(opsi_key, ('F8FAFC','374151'))[1]
    write_cell(row.cells[0], no_label if is_first else '', size=9, bold=True,
               color=hx(HEADER_BG), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg='F1F5F9', v='top', mar=(80,60,60,60))
    sc = row.cells[1]; shd(sc, SOAL_BG); valign(sc, 'top'); cell_mar(sc, 60, 60, 100, 80)
    for p in sc.paragraphs[1:]: p._p.getparent().remove(p._p)
    p0 = sc.paragraphs[0]; p0.clear(); pspacing(p0, 0, 0)
    run_text(p0, pasangan_kiri if is_first else '', 'Calibri', 9.5)
    jc = row.cells[2]; shd(jc, 'E0F2FE'); valign(jc, 'center'); cell_mar(jc, 60, 60, 60, 60)
    for p in jc.paragraphs[1:]: p._p.getparent().remove(p._p)
    pj = jc.paragraphs[0]; pj.clear(); pj.alignment = WD_ALIGN_PARAGRAPH.CENTER; pspacing(pj, 0, 0)
    run_text(pj, 'JODOH' if is_first else '', 'Calibri', 8, True, False, hx(HEADER_BG))
    write_cell(row.cells[3], opsi_key, size=9, bold=True,
               color=hx(fg_clr), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg=row_bg, v='center', mar=(60,60,40,40))
    write_cell(row.cells[4], pasangan_kanan, size=9,
               color=hx(fg_clr), bg=row_bg, v='top', mar=(60,60,100,80))
    write_cell(row.cells[5], opsi_key, size=9, bold=True,
               color=hx(CORRECT_FG), align=WD_ALIGN_PARAGRAPH.CENTER,
               bg=CORRECT_BG, v='center', mar=(60,60,40,40))
    return row

def add_separator(table, label):
    """Baris pemisah antar jenis soal."""
    row = table.add_row()
    set_row_widths(row)
    for ci in range(6):
        shd(row.cells[ci], HEADER_BG)
        cell_mar(row.cells[ci], 50, 50, 80, 80)
    sc = row.cells[0]
    for p in sc.paragraphs[1:]: p._p.getparent().remove(p._p)
    # Merge semua cells jadi satu
    merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2]).merge(
        row.cells[3]).merge(row.cells[4]).merge(row.cells[5])
    shd(merged, HEADER_BG); cell_mar(merged, 60, 60, 120, 80)
    for p in merged.paragraphs[1:]: p._p.getparent().remove(p._p)
    pm = merged.paragraphs[0]; pm.clear(); pm.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pspacing(pm, 0, 0)
    run_text(pm, f'  {label}', 'Calibri', 9, True, False, hx('FFFFFF'))

# ════════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════════════
doc = Document()
# Hapus paragraf kosong default
for p in doc.paragraphs: p._p.getparent().remove(p._p)

sec = doc.sections[0]
sec.page_width   = Cm(21)
sec.page_height  = Cm(29.7)
sec.left_margin  = sec.right_margin = Cm(1.5)
sec.top_margin   = sec.bottom_margin = Cm(1.5)

# ── JUDUL DOKUMEN ──────────────────────────────────────────────────────────────
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
pspacing(tp, 0, 60)
run_text(tp, 'TEMPLATE BANK SOAL — CBT SCHOOL ENTERPRISE', 'Calibri', 14, True, False, hx(HEADER_BG))
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
pspacing(sp, 0, 120)
run_text(sp, 'Isi tabel di bawah sesuai contoh  |  Setiap baris OPSI = satu pilihan jawaban  |  Kolom KUNCI diisi V pada jawaban yang benar',
         'Calibri', 8.5, False, True, hx('64748B'))

# ── PANDUAN SINGKAT ────────────────────────────────────────────────────────────
pg = doc.add_table(rows=1, cols=5)
pg.alignment = WD_TABLE_ALIGNMENT.CENTER
no_border_table(pg)
tbl_width(pg, TOTAL_W)
for ci in range(5):
    cell_width(pg.rows[0].cells[ci], TOTAL_W//5)

guide = [
    ('PG BIASA',  '2563EB', 'DBEAFE', 'Pilihan ganda biasa.\nOpsi A–E, KUNCI = V\npada jawaban benar.'),
    ('PG KOMPLEKS','7C3AED','EDE9FE', 'Pilihan ganda\nbisa lebih dari\n1 jawaban benar.'),
    ('BENAR/SALAH','065F46','D1FAE5', 'Pilih BENAR atau\nSALAH. Isi KUNCI = V\npada jawaban yang benar.'),
    ('MENJODOHKAN','B45309','FEF9C3', 'Pasangkan kiri-kanan.\nKolom KUNCI diisi\nhuruf pasangannya.'),
    ('ESAI',       '991B1B','FEE2E2', 'Soal uraian/esai.\nIsi rubrik penilaian\ndi kolom JAWABAN.'),
]
for ci, (lbl, fg, bg, desc) in enumerate(guide):
    gc = pg.rows[0].cells[ci]
    shd(gc, bg); cell_mar(gc, 80, 80, 80, 80); valign(gc, 'top')
    for p in gc.paragraphs[1:]: p._p.getparent().remove(p._p)
    p0 = gc.paragraphs[0]; p0.clear(); p0.alignment = WD_ALIGN_PARAGRAPH.CENTER; pspacing(p0, 0, 40)
    run_text(p0, lbl, 'Calibri', 8.5, True, False, hx(fg))
    p1 = gc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER; pspacing(p1, 0, 0)
    run_text(p1, desc, 'Calibri', 7.5, False, False, hx('374151'))

sp2 = doc.add_paragraph(); pspacing(sp2, 80, 0)

# ── TABEL UTAMA ────────────────────────────────────────────────────────────────
t = make_main_table(doc)
add_header_row(t)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SOAL 1 — PG BIASA (dengan gambar)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_separator(t, '1 — PG BIASA (Pilihan Ganda)  |  Isi KUNCI = V pada satu jawaban yang benar')
img1 = make_photo_placeholder(320, 160, 'Gambar Soal Pertanyaan 1')
add_soal_row(t, 1,
    'Perhatikan gambar di samping. Proses yang ditunjukkan gambar tersebut disebut ...',
    'PG', '93C5FD', with_image=True, image_buf=img1)
options1 = [('A','Fotosintesis',False),('B','Respirasi',False),('C','Transpirasi',True),('D','Fermentasi',False),('E','Katabolisme',False)]
for opsi, jwb, correct in options1:
    add_option_row(t, 1, opsi, jwb, 'V' if correct else '', is_correct=correct)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SOAL 2 — PG BIASA (tanpa gambar)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_soal_row(t, 2,
    'Negara Indonesia memproklamasikan kemerdekaannya pada tanggal ...',
    'PG', '93C5FD')
options2 = [('A','17 Agustus 1945',True),('B','1 Juni 1945',False),('C','28 Oktober 1928',False),('D','22 Juni 1945',False)]
for opsi, jwb, correct in options2:
    add_option_row(t, 2, opsi, jwb, 'V' if correct else '', is_correct=correct)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SOAL 3 — PG KOMPLEKS (lebih dari 1 jawaban benar)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_separator(t, '2 — PG KOMPLEKS (Pilihan Ganda)  |  Bisa lebih dari 1 jawaban benar  |  KUNCI = V pada semua yang benar')
add_soal_row(t, 3,
    'Manakah pernyataan berikut yang BENAR mengenai sel? (Pilih semua yang benar)',
    'PG+', 'C4B5FD')
options3 = [('A','Sel adalah unit terkecil kehidupan',True),('B','Semua sel memiliki dinding sel',False),('C','Sel memiliki membran plasma',True),('D','Virus tergolong sel prokariot',False),('E','Mitokondria adalah organel sel eukariot',True)]
for opsi, jwb, correct in options3:
    add_option_row(t, 3, opsi, jwb, 'V' if correct else '', is_correct=correct)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SOAL 4 — BENAR / SALAH
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_separator(t, '3 — BENAR / SALAH  |  Isi KUNCI = V pada kolom BENAR atau SALAH yang benar')
add_tf_row(t, 4, 'Matahari terbit dari arah timur.', 'B')
add_tf_row(t, 5, 'Bumi berputar mengelilingi bulan.', 'S')
add_tf_row(t, 6, 'Air mendidih pada suhu 100°C pada tekanan normal.', 'B')

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SOAL 5 — MENJODOHKAN
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_separator(t, '4 — MENJODOHKAN  |  Kolom KUNCI diisi huruf pasangan yang sesuai (A, B, C, ...)')
pasangan = [
    ('A', 'Soekarno',      'Presiden pertama RI'),
    ('B', 'Ki Hajar Dewantara', 'Bapak Pendidikan Nasional'),
    ('C', 'R.A. Kartini',  'Pahlawan emansipasi wanita'),
    ('D', 'Imam Bonjol',   'Pahlawan dari Sumatera Barat'),
]
for i, (opsi, kiri, kanan) in enumerate(pasangan):
    add_matching_pair(t, str(7), kiri, kanan, opsi, is_first=(i==0))

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SOAL 6 — ESAI
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
add_separator(t, '5 — ESAI / URAIAN  |  Kolom JAWABAN = rubrik/kunci jawaban. Kolom KUNCI = skor maksimum (angka)')
add_essay_row(t, 8,
    'Jelaskan pengertian Pancasila dan sebutkan 5 silanya!',
    'ESAI', 'FEF3C7',
    'Pancasila adalah dasar negara Indonesia yang terdiri dari 5 sila:\n'
    '1. Ketuhanan YME\n2. Kemanusiaan\n3. Persatuan\n4. Kerakyatan\n5. Keadilan Sosial.\n'
    '(Skor: 2 poin per sila + 2 poin definisi)')
add_essay_row(t, 9,
    'Hitunglah luas persegi panjang dengan panjang 12 cm dan lebar 8 cm, serta jelaskan caranya!',
    'ESAI', 'FEF3C7',
    'L = p × l = 12 × 8 = 96 cm².\nRubrik: rumus benar (4), perhitungan benar (4), satuan benar (2).')

# ─── CATATAN BAWAH ────────────────────────────────────────────────────────────
fn = doc.add_paragraph()
pspacing(fn, 120, 0)
fn.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_text(fn,
    'CATATAN:  Kolom JENIS otomatis terdeteksi saat import. '
    'Nilai di kolom KUNCI: V = benar (PG), huruf pasangan (Menjodohkan), angka skor (Esai). '
    'Hapus baris CONTOH di atas dan ganti dengan soal Anda sendiri.',
    'Calibri', 8, False, True, hx('64748B'))

# ─── SIMPAN ───────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f'[OK] Template disimpan: {OUTPUT_PATH}')
print(f'     Ukuran: {os.path.getsize(OUTPUT_PATH):,} bytes')
